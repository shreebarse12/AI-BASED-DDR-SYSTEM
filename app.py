"""
DDR Report Generator — Pure Streamlit App
No FastAPI, no backend server. Everything runs in Streamlit.
"""

import os
import io
import re
import json
import uuid
import base64
import tempfile
from datetime import datetime
from pathlib import Path

import streamlit as st
from groq import Groq
import pdfplumber
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, Image as RLImage
)
from reportlab.lib.enums import TA_LEFT
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# PyMuPDF — optional
try:
    import fitz
    FITZ_AVAILABLE = True
except Exception:
    FITZ_AVAILABLE = False

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="DDR Report Generator",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
  /* Force light background regardless of user theme */
  html, body, [data-testid="stAppViewContainer"], [data-testid="stApp"] {
      background-color: #f0f2f5 !important;
      color: #1a1a1a !important;
  }
  [data-testid="stHeader"] { background: #1d3557 !important; }
  .block-container {
      max-width: 860px !important;
      padding-top: 2rem !important;
      background: transparent;
  }

  /* Header */
  .header-bar {
      background: #1d3557;
      padding: 16px 28px;
      border-radius: 10px;
      margin-bottom: 28px;
      display: flex;
      align-items: center;
      gap: 14px;
      box-shadow: 0 2px 8px rgba(29,53,87,0.25);
  }
  .header-bar .badge {
      background: #e63946; color: white; font-weight: 700;
      font-size: 13px; padding: 5px 12px; border-radius: 5px;
      letter-spacing: 0.5px;
  }
  .header-bar .title { color: white; font-size: 19px; font-weight: 600; }
  .header-bar .sub   { color: rgba(255,255,255,0.55); font-size: 13px; }

  /* Upload section */
  .upload-card {
      background: white;
      border-radius: 10px;
      padding: 24px 28px;
      box-shadow: 0 1px 4px rgba(0,0,0,0.08);
      margin-bottom: 20px;
  }
  .upload-card h3 {
      font-size: 15px; font-weight: 600; color: #1d3557;
      margin-bottom: 16px; padding-bottom: 10px;
      border-bottom: 2px solid #e8ecf0;
  }

  /* Report card */
  .report-card {
      background: white;
      border-radius: 10px;
      padding: 28px 32px;
      box-shadow: 0 1px 4px rgba(0,0,0,0.08);
      margin-top: 20px;
  }

  /* Section headers */
  .sec-heading {
      display: flex; align-items: center; gap: 10px;
      border-bottom: 2px solid #e8ecf0;
      padding-bottom: 8px; margin: 24px 0 14px 0;
  }
  .sec-num {
      background: #1d3557; color: white; font-size: 10px; font-weight: 700;
      padding: 3px 8px; border-radius: 4px; letter-spacing: 0.5px;
  }
  .sec-title { font-size: 15px; font-weight: 600; color: #1d3557; }

  /* Area cards */
  .area-card {
      border: 1px solid #e4e8ed; border-radius: 8px;
      padding: 16px 18px; margin-bottom: 12px;
      background: #fafbfc;
  }
  .area-name { font-weight: 600; font-size: 14px; margin-bottom: 7px; color: #111; }
  .area-obs  { font-size: 13px; color: #444; line-height: 1.7; }

  /* Severity badges */
  .badge-high   { background:#fde8e8; color:#c0392b; padding:3px 9px; border-radius:4px; font-size:11px; font-weight:600; border:1px solid #fecaca; }
  .badge-medium { background:#fef3cd; color:#856404; padding:3px 9px; border-radius:4px; font-size:11px; font-weight:600; border:1px solid #fde68a; }
  .badge-low    { background:#d4edda; color:#155724; padding:3px 9px; border-radius:4px; font-size:11px; font-weight:600; border:1px solid #bbf7d0; }

  /* Thermal row */
  .thermal-row {
      margin-top:10px; padding:9px 14px; background:#eef4fb;
      border-left:3px solid #457b9d; border-radius:0 5px 5px 0;
      font-size:13px; color:#1e3a5f; line-height:1.5;
  }

  /* Tags */
  .tag { display:inline-block; font-size:12px; background:#f0f2f5; border:1px solid #d8dde4; border-radius:4px; padding:4px 11px; color:#555; margin:3px; }
  .tag-conflict { background:#fff5f5; border-color:#f5c6c6; color:#c0392b; }

  /* Streamlit overrides */
  .stFileUploader > div { background: white !important; border-radius: 8px !important; border: 1px solid #e4e8ed !important; }
  div[data-testid="stFileUploaderDropzone"] { background: #fafbfc !important; }
  .stButton > button {
      border-radius: 6px !important; font-weight: 600 !important;
      padding: 10px 24px !important; font-size: 14px !important;
  }
  .stDownloadButton > button {
      background: white !important; color: #1d3557 !important;
      border: 1.5px solid #c8d0dc !important; border-radius: 6px !important;
      font-weight: 600 !important; padding: 8px 20px !important;
  }
  .stDownloadButton > button:hover {
      background: #f4f6f9 !important; border-color: #1d3557 !important;
  }
  .stSuccess { background: #f0fdf4 !important; color: #166534 !important; border-radius: 8px !important; }
  .stAlert   { border-radius: 8px !important; }
  [data-testid="stDataFrame"] { border-radius: 8px; overflow: hidden; }
  label[data-testid="stWidgetLabel"] { color: #333 !important; font-weight: 600 !important; font-size: 13px !important; }
  p, li, div { color: #222 !important; }
  h1,h2,h3 { color: #1d3557 !important; }
</style>
""", unsafe_allow_html=True)

# ── Header ─────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="header-bar">
  <span class="badge">DDR</span>
  <span class="title">Report Generator</span>
  <span class="sub">AI-Assisted Diagnostic Reports</span>
</div>
""", unsafe_allow_html=True)

# ── Groq client ────────────────────────────────────────────────────────────────
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
if not GROQ_API_KEY:
    GROQ_API_KEY = st.secrets.get("GROQ_API_KEY", "") if hasattr(st, "secrets") else ""

client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

SYSTEM_PROMPT = """You are a senior building diagnostics engineer with 20+ years of experience.
Your task: read raw inspection and thermal report data, then produce a structured DDR (Detailed Diagnostic Report) as JSON.

STRICT RULES:
- NEVER invent facts not in the documents
- If data conflicts between documents → explicitly note the conflict
- If data is missing → use exactly "Not Available"
- Use plain, client-friendly language
- Be specific: mention exact locations, temperatures, measurements when available

Return ONLY valid JSON with this exact structure:
{
  "property_summary": "2-4 sentence executive overview of all key issues",
  "areas": [
    {
      "name": "Exact area name",
      "severity": "High | Medium | Low",
      "observations": "Detailed observation combining inspection + thermal findings",
      "thermal_finding": "Specific temperature data or thermal anomaly, or 'Not Available'",
      "image_label": "Short label or null"
    }
  ],
  "root_causes": [{"issue": "Issue title", "cause": "Probable root cause with reasoning"}],
  "severity_assessment": [{"area": "Area name", "severity": "High | Medium | Low", "reasoning": "Why this severity"}],
  "recommended_actions": ["Specific, actionable recommendation"],
  "additional_notes": "Important context or 'Not Available'",
  "missing_info": ["List of missing data points"],
  "conflicts": ["List of conflicting data or empty array"]
}"""

# ── Core functions ─────────────────────────────────────────────────────────────

def extract_text_from_pdf(data: bytes) -> str:
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = []
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        for row in table:
                            if row:
                                text += "\n" + " | ".join(str(c) for c in row if c)
                pages.append(f"[Page {i+1}]\n{text}")
            return "\n\n".join(pages)
    except Exception as e:
        try:
            reader = PdfReader(io.BytesIO(data))
            return "\n\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception:
            return f"[Could not extract text: {e}]"


def extract_text(data: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(data)
    elif ext in ("txt", "md"):
        return data.decode("utf-8", errors="replace")
    elif ext in ("png", "jpg", "jpeg"):
        return "[IMAGE FILE — visual content only]"
    return data.decode("utf-8", errors="replace")


def extract_images_from_pdf(data: bytes, label: str, max_pages: int = 4) -> list:
    images = []
    if FITZ_AVAILABLE:
        try:
            doc = fitz.open(stream=data, filetype="pdf")
            for i in range(min(len(doc), max_pages)):
                page = doc[i]
                pix = page.get_pixmap(matrix=fitz.Matrix(1.2, 1.2))
                b64 = base64.b64encode(pix.tobytes("jpeg")).decode()
                images.append({"src": f"data:image/jpeg;base64,{b64}", "label": f"{label} — Page {i+1}"})
            doc.close()
            return images
        except Exception:
            pass
    try:
        from pdf2image import convert_from_bytes
        pages = convert_from_bytes(data, dpi=100, first_page=1, last_page=max_pages)
        for i, pg in enumerate(pages):
            buf = io.BytesIO()
            pg.save(buf, format="JPEG", quality=70)
            b64 = base64.b64encode(buf.getvalue()).decode()
            images.append({"src": f"data:image/jpeg;base64,{b64}", "label": f"{label} — Page {i+1}"})
    except Exception:
        pass
    return images


def call_groq(inspection_text: str, thermal_text: str) -> dict:
    user_msg = f"""INSPECTION REPORT:
{inspection_text[:6000]}

THERMAL REPORT:
{thermal_text[:4000]}

Generate the DDR JSON now. Return ONLY valid JSON, no preamble, no markdown fences."""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        max_tokens=2000,
        temperature=0.1,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": user_msg}
        ]
    )
    raw = response.choices[0].message.content or ""
    clean = raw.strip()
    if clean.startswith("```"):
        clean = clean.split("```")[1]
        if clean.startswith("json"):
            clean = clean[4:]
    clean = clean.strip().rstrip("`").strip()
    try:
        return json.loads(clean)
    except json.JSONDecodeError:
        match = re.search(r'\{[\s\S]*\}', clean)
        if match:
            return json.loads(match.group())
        raise ValueError("Could not parse AI response as JSON. Raw: " + raw[:300])


def build_pdf(report: dict, images: list) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm, topMargin=2.5*cm, bottomMargin=2*cm)
    W = A4[0] - 4*cm
    styles = getSampleStyleSheet()
    DARK   = colors.HexColor("#1d3557")
    ACCENT = colors.HexColor("#e63946")
    MUTED  = colors.HexColor("#6b7280")

    title_s = ParagraphStyle("T", parent=styles["Title"], fontSize=20,
        textColor=DARK, spaceAfter=4, fontName="Helvetica-Bold")
    sub_s   = ParagraphStyle("S", parent=styles["Normal"], fontSize=9,
        textColor=MUTED, spaceAfter=16)
    h_s     = ParagraphStyle("H", parent=styles["Heading1"], fontSize=12,
        textColor=DARK, fontName="Helvetica-Bold", spaceBefore=14, spaceAfter=6)
    body_s  = ParagraphStyle("B", parent=styles["Normal"], fontSize=10,
        textColor=colors.HexColor("#374151"), leading=15, spaceAfter=5)
    lbl_s   = ParagraphStyle("L", parent=styles["Normal"], fontSize=8,
        textColor=MUTED, fontName="Helvetica-Oblique")

    def sev_clr(s):
        s = (s or "").lower()
        if s == "high":   return colors.HexColor("#dc2626")
        if s == "medium": return colors.HexColor("#d97706")
        return colors.HexColor("#16a34a")

    def sec(num, title):
        return [HRFlowable(width=W, thickness=2, color=ACCENT, spaceAfter=4),
                Paragraph(f"<font color='#e63946'>{num}</font>  {title}", h_s)]

    story = []
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph("DETAILED DIAGNOSTIC REPORT", title_s))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}  ·  AI-Synthesised", sub_s))
    story.append(HRFlowable(width=W, thickness=3, color=ACCENT, spaceAfter=16))

    # 01
    story += sec("01", "Property Issue Summary")
    story.append(Paragraph(report.get("property_summary", "Not Available"), body_s))
    story.append(Spacer(1, 0.2*cm))

    # 02
    story += sec("02", "Area-wise Observations")
    for area in report.get("areas", []):
        sev = area.get("severity", "Medium")
        sc  = sev_clr(sev)
        row1 = Paragraph(f'<b>{area.get("name","")}</b>  <font color="{sc.hexval()}" size="8">[{sev.upper()}]</font>', body_s)
        row2 = Paragraph(area.get("observations", ""), body_s)
        tf   = area.get("thermal_finding", "Not Available")
        row3 = Paragraph(f'<font color="#2563eb"><b>Thermal:</b> {tf}</font>', body_s)
        tbl  = Table([[row1],[row2],[row3]], colWidths=[W])
        tbl.setStyle(TableStyle([
            ("BOX",(0,0),(-1,-1),0.5,colors.HexColor("#d1d5db")),
            ("BACKGROUND",(0,0),(0,0),colors.HexColor("#f9fafb")),
            ("LEFTPADDING",(0,0),(-1,-1),10), ("RIGHTPADDING",(0,0),(-1,-1),10),
            ("TOPPADDING",(0,0),(-1,-1),7),   ("BOTTOMPADDING",(0,0),(-1,-1),7),
        ]))
        story.append(tbl)
        if images:
            try:
                b64 = images[0]["src"].split(",",1)[1]
                rl  = RLImage(io.BytesIO(base64.b64decode(b64)), width=W*0.55, height=W*0.32)
                story.append(rl)
                story.append(Paragraph(images[0]["label"], lbl_s))
            except Exception:
                pass
        story.append(Spacer(1, 0.2*cm))

    # 03
    story += sec("03", "Probable Root Cause")
    for i, rc in enumerate(report.get("root_causes",[]), 1):
        story.append(Paragraph(f"<b>{i}. {rc.get('issue','')}</b>", body_s))
        story.append(Paragraph(rc.get("cause",""), body_s))

    # 04
    story += sec("04", "Severity Assessment")
    rows = [["Area","Severity","Reasoning"]]
    for s in report.get("severity_assessment",[]):
        rows.append([s.get("area",""), s.get("severity",""), s.get("reasoning","")])
    if len(rows) > 1:
        t = Table(rows, colWidths=[W*0.25, W*0.12, W*0.63])
        t.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,0),DARK), ("TEXTCOLOR",(0,0),(-1,0),colors.white),
            ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"), ("FONTSIZE",(0,0),(-1,-1),9),
            ("GRID",(0,0),(-1,-1),0.5,colors.HexColor("#e5e7eb")),
            ("LEFTPADDING",(0,0),(-1,-1),8), ("RIGHTPADDING",(0,0),(-1,-1),8),
            ("TOPPADDING",(0,0),(-1,-1),5),  ("BOTTOMPADDING",(0,0),(-1,-1),5),
            ("VALIGN",(0,0),(-1,-1),"TOP"),
        ]))
        story.append(t)
    story.append(Spacer(1, 0.2*cm))

    # 05
    story += sec("05", "Recommended Actions")
    for i, a in enumerate(report.get("recommended_actions",[]), 1):
        story.append(Paragraph(f"{i}.  {a}", body_s))

    # 06
    story += sec("06", "Additional Notes")
    story.append(Paragraph(report.get("additional_notes","Not Available"), body_s))

    # 07
    story += sec("07", "Missing or Unclear Information")
    missing   = report.get("missing_info", [])
    conflicts = report.get("conflicts", [])
    if not missing and not conflicts:
        story.append(Paragraph("No missing information identified.", body_s))
    for m in missing:   story.append(Paragraph(f"⚠  {m}", body_s))
    for c in conflicts: story.append(Paragraph(f"⚡  CONFLICT: {c}", body_s))

    doc.build(story)
    return buf.getvalue()


def build_docx(report: dict, images: list) -> bytes:
    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    title = doc.add_heading("Detailed Diagnostic Report", 0)
    title.runs[0].font.color.rgb = RGBColor(0x1d, 0x35, 0x57)
    title.runs[0].font.size = Pt(22)

    meta = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}  ·  AI-Synthesised")
    meta.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    meta.runs[0].font.size = Pt(9)

    def add_hr():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),'e63946')
        pBdr.append(bot); pPr.append(pBdr)

    def add_sec(num, txt):
        add_hr()
        h = doc.add_heading(f"{num}  {txt}", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1d, 0x35, 0x57)
        h.runs[0].font.size = Pt(13)

    def sev_rgb(s):
        s = (s or "").lower()
        if s == "high":   return RGBColor(0xdc, 0x26, 0x26)
        if s == "medium": return RGBColor(0xd9, 0x77, 0x06)
        return RGBColor(0x16, 0xa3, 0x4a)

    add_sec("01", "Property Issue Summary")
    doc.add_paragraph(report.get("property_summary","Not Available"))

    add_sec("02", "Area-wise Observations")
    for area in report.get("areas",[]):
        sev = area.get("severity","Medium")
        p   = doc.add_paragraph()
        r   = p.add_run(area.get("name","")); r.bold = True; r.font.size = Pt(11)
        sr  = p.add_run(f"  [{sev.upper()}]")
        sr.font.color.rgb = sev_rgb(sev); sr.font.size = Pt(9)
        doc.add_paragraph(area.get("observations",""))
        tp = doc.add_paragraph()
        tr = tp.add_run(f"Thermal: {area.get('thermal_finding','Not Available')}")
        tr.font.color.rgb = RGBColor(0x25, 0x63, 0xeb); tr.font.size = Pt(9)
        if images:
            try:
                b64 = images[0]["src"].split(",",1)[1]
                buf = io.BytesIO(base64.b64decode(b64))
                doc.add_picture(buf, width=Inches(4))
            except Exception:
                doc.add_paragraph("[Image Not Available]")

    add_sec("03", "Probable Root Cause")
    for i, rc in enumerate(report.get("root_causes",[]), 1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(f"{rc.get('issue','')}: "); r.bold = True
        p.add_run(rc.get("cause",""))

    add_sec("04", "Severity Assessment")
    sev_data = report.get("severity_assessment",[])
    if sev_data:
        tbl = doc.add_table(rows=1, cols=3); tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, txt in enumerate(["Area","Severity","Reasoning"]):
            hdr[i].text = txt
            hdr[i].paragraphs[0].runs[0].bold = True
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff,0xff,0xff)
            sh = OxmlElement('w:shd')
            sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),'1d3557')
            hdr[i]._tc.get_or_add_tcPr().append(sh)
        for s in sev_data:
            row = tbl.add_row().cells
            row[0].text = s.get("area","")
            row[1].text = s.get("severity","")
            if row[1].paragraphs[0].runs:
                row[1].paragraphs[0].runs[0].font.color.rgb = sev_rgb(s.get("severity",""))
            row[2].text = s.get("reasoning","")

    add_sec("05", "Recommended Actions")
    for a in report.get("recommended_actions",[]):
        doc.add_paragraph(a, style="List Number")

    add_sec("06", "Additional Notes")
    doc.add_paragraph(report.get("additional_notes","Not Available"))

    add_sec("07", "Missing or Unclear Information")
    missing = report.get("missing_info",[]); conflicts = report.get("conflicts",[])
    if not missing and not conflicts:
        doc.add_paragraph("No missing information identified.")
    for m in missing:
        doc.add_paragraph(m, style="List Bullet")
    for c in conflicts:
        p = doc.add_paragraph(); r = p.add_run(f"CONFLICT: {c}")
        r.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)

    out = io.BytesIO(); doc.save(out); return out.getvalue()


# ── UI ─────────────────────────────────────────────────────────────────────────

def sec_header(num, title):
    st.markdown(f"""
    <div class="sec-heading">
      <span class="sec-num">{num}</span>
      <span class="sec-title">{title}</span>
    </div>""", unsafe_allow_html=True)

def sev_badge(s):
    s = (s or "medium").lower()
    cls = "badge-high" if s=="high" else "badge-low" if s=="low" else "badge-medium"
    return f'<span class="{cls}">{s.capitalize()}</span>'


# API key check
if not GROQ_API_KEY:
    st.warning("⚠️ GROQ_API_KEY not set. Add it to your environment or Streamlit secrets.")
    st.code('$env:GROQ_API_KEY="gsk_your_key_here"  # PowerShell\nexport GROQ_API_KEY=gsk_your_key_here  # Mac/Linux')
    st.stop()

st.markdown('<div class="upload-card">', unsafe_allow_html=True)
st.markdown('<h3>📂 Upload Inspection Documents</h3>', unsafe_allow_html=True)
col1, col2 = st.columns(2)
with col1:
    insp_file = st.file_uploader("📋 Inspection Report", type=["pdf","txt","png","jpg","jpeg"],
        help="Upload your site inspection report (PDF, TXT, or image)")
with col2:
    therm_file = st.file_uploader("🌡️ Thermal Report", type=["pdf","txt","png","jpg","jpeg"],
        help="Upload your thermal imaging report (PDF, TXT, or image)")
st.markdown('</div>', unsafe_allow_html=True)
st.markdown("")

if st.button("🔍 Generate DDR Report", type="primary", disabled=not (insp_file and therm_file)):

    prog = st.progress(0, text="Reading documents...")
    try:
        insp_data  = insp_file.read()
        therm_data = therm_file.read()

        prog.progress(20, text="Extracting text...")
        insp_text  = extract_text(insp_data,  insp_file.name)
        therm_text = extract_text(therm_data, therm_file.name)

        prog.progress(35, text="Extracting images...")
        images = []
        if insp_file.name.lower().endswith(".pdf"):
            images += extract_images_from_pdf(insp_data, "Inspection")
        elif insp_file.name.lower().endswith((".png",".jpg",".jpeg")):
            b64 = base64.b64encode(insp_data).decode()
            images.append({"src": f"data:image/jpeg;base64,{b64}", "label": "Inspection Image"})

        if therm_file.name.lower().endswith(".pdf"):
            images += extract_images_from_pdf(therm_data, "Thermal")
        elif therm_file.name.lower().endswith((".png",".jpg",".jpeg")):
            b64 = base64.b64encode(therm_data).decode()
            images.append({"src": f"data:image/jpeg;base64,{b64}", "label": "Thermal Image"})

        prog.progress(55, text="Running AI analysis...")
        report = call_groq(insp_text, therm_text)

        prog.progress(90, text="Building report...")
        st.session_state["report"] = report
        st.session_state["images"] = images
        st.session_state["report_id"] = str(uuid.uuid4())[:8]
        st.session_state["generated_at"] = datetime.now().strftime("%d %B %Y, %H:%M")
        prog.progress(100, text="Done!")

    except Exception as e:
        prog.empty()
        st.error(f"⚠️ {str(e)}")
        st.stop()

    prog.empty()
    st.success("✅ Report generated!")


# ── Render report ──────────────────────────────────────────────────────────────
if "report" in st.session_state:
    report  = st.session_state["report"]
    images  = st.session_state["images"]
    rid     = st.session_state["report_id"]
    gen_at  = st.session_state["generated_at"]

    st.markdown('<div class="report-card">', unsafe_allow_html=True)

    # Report meta + export
    mc1, mc2, mc3 = st.columns([3, 1.1, 1.1])
    with mc1:
        st.markdown(f"<p style='color:#888;font-size:13px;margin:0'><b>Report ID:</b> <code>{rid}</code> &nbsp;·&nbsp; <b>Generated:</b> {gen_at}</p>", unsafe_allow_html=True)
    with mc2:
        try:
            pdf_bytes = build_pdf(report, images)
            st.download_button("⬇ Export PDF", pdf_bytes,
                file_name=f"DDR_Report_{rid}.pdf", mime="application/pdf", use_container_width=True)
        except Exception as e:
            st.error(f"PDF error: {e}")
    with mc3:
        try:
            docx_bytes = build_docx(report, images)
            st.download_button("⬇ Export DOCX", docx_bytes,
                file_name=f"DDR_Report_{rid}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True)
        except Exception as e:
            st.error(f"DOCX error: {e}")

    st.markdown("<hr style='border:none;border-top:2px solid #e8ecf0;margin:16px 0'>", unsafe_allow_html=True)

    # 01 Summary
    sec_header("01", "Property Issue Summary")
    st.write(report.get("property_summary", "Not Available"))

    # 02 Areas
    sec_header("02", "Area-wise Observations")
    img_idx = 0
    for area in report.get("areas", []):
        sev     = area.get("severity", "Medium")
        thermal = area.get("thermal_finding", "")
        thermal_html = f'<div class="thermal-row"><b>Thermal:</b> {thermal}</div>' if thermal and thermal != "Not Available" else ""
        st.markdown(f"""
        <div class="area-card">
          <div style="font-weight:600;font-size:14px;margin-bottom:6px">
            {area.get('name','')} &nbsp; {sev_badge(sev)}
          </div>
          <div style="font-size:13px;color:#444;line-height:1.65">{area.get('observations','')}</div>
          {thermal_html}
        </div>""", unsafe_allow_html=True)
        if images and img_idx < len(images):
            img = images[img_idx]; img_idx += 1
            st.image(img["src"], caption=img.get("label",""), width=380)

    # 03 Root Cause
    sec_header("03", "Probable Root Cause")
    for i, rc in enumerate(report.get("root_causes", []), 1):
        st.markdown(f"**{i}. {rc.get('issue','')}:** {rc.get('cause','')}")

    # 04 Severity Assessment
    sec_header("04", "Severity Assessment")
    sev_data = report.get("severity_assessment", [])
    if sev_data:
        import pandas as pd
        df = pd.DataFrame([{"Area": s.get("area",""), "Severity": s.get("severity",""), "Reasoning": s.get("reasoning","")} for s in sev_data])
        st.dataframe(df, use_container_width=True, hide_index=True)

    # 05 Recommended Actions
    sec_header("05", "Recommended Actions")
    for i, a in enumerate(report.get("recommended_actions", []), 1):
        st.markdown(f"{i}. {a}")

    # 06 Additional Notes
    sec_header("06", "Additional Notes")
    st.write(report.get("additional_notes", "Not Available"))

    # 07 Missing Info
    sec_header("07", "Missing or Unclear Information")
    missing   = report.get("missing_info", [])
    conflicts = report.get("conflicts", [])
    if not missing and not conflicts:
        st.write("No missing information identified.")
    else:
        tags = "".join(f'<span class="tag">⚠ {m}</span>' for m in missing)
        tags += "".join(f'<span class="tag tag-conflict">⚡ Conflict: {c}</span>' for c in conflicts)
        st.markdown(tags, unsafe_allow_html=True)

    # Gallery
    if images:
        sec_header("A", "Extracted Document Images")
        cols = st.columns(3)
        for i, img in enumerate(images):
            with cols[i % 3]:
                st.image(img["src"], caption=img.get("label",""), use_container_width=True)

    st.markdown('</div>', unsafe_allow_html=True)